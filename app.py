import streamlit as st
from io import BytesIO
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Sözleşme Oluşturucu", page_icon="🏠", layout="centered")
st.title("🏠Emlak Sözleşme Asistanı")
st.caption("Kira ve alım-satım sözleşmelerini kolayca hazırlayın ve indirin.")

def create_sales_contract(cins, il, ilce, mahalle, sokak, kapi_no,
                          pafta_no, ada_no, parsel_no, alici, satici,
                          fiyat, kaparo, odeme_sekli, emlakci_komisyon_yuzdesi,
                          mahkemeleri, date):
    doc = docx.Document()
    heading = doc.add_heading('EMLAK ALIM SATIM SÖZLEŞMESİ', 0)
    run = heading.runs[0]
    run.font.name = 'Time New Roman'
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.bold = True
    run.font.size = Pt(16)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'CİNSİ   : {cins}')
    doc.add_paragraph(f'İL      : {il}')
    doc.add_paragraph(f'İLÇE    : {ilce}')
    doc.add_paragraph(f'MAHALLE : {mahalle}')
    doc.add_paragraph(f'SOKAK   : {sokak}')
    doc.add_paragraph(f'KAPI NO : {kapi_no}')
    doc.add_paragraph(f'PAFTA NO: {pafta_no}')
    doc.add_paragraph(f'ADA NO  : {ada_no}')
    doc.add_paragraph(f'PARSEL NO:{parsel_no}')
    doc.add_paragraph(f'ALICI {alici} ile SATICI {satici} yukarıda kayıtlı emlağın aşağıda belirlenen koşullarda satılması için anlaşmışlardır.')
    doc.add_paragraph(f"1- SATICI, sahibi bulunduğu yukarıda kayıtlı emlağını {fiyat} TL'dan satmayı kabul etmiştir.")
    doc.add_paragraph(f"2- Satış bedeline mahsuben, ALICI’dan {kaparo} TL kaparo alınmıştır.")
    doc.add_paragraph(f"3- Satış bedelinin ödeme şekli {odeme_sekli}")
    doc.add_paragraph("4- Bu sözleşme imzalandıktan sonra, borçlar kanunu’nun 156/2 maddesine göre, taraflardan ALICI, bu emlağı almaktan vazgeçtiği taktirde, verdiği kaporayı geri alamayacaktır. SATICI bu emlağı satmaktan vazgeçerse, kaporayı iade edecek ve kapora miktarı kadar daha tazminat ödeyecektir.")
    doc.add_paragraph(f"5- ALICI ve SATICI, kendilerine bu sözleşmeyi sağlayan EMLAK KOMİSYONCUSU’na sözleşmenin imzasından itibaren, gerçek satış bedeli üzerinden %{emlakci_komisyon_yuzdesi} +KDV tutarında hizmet ücreti ödemeyi kabul ederler.")
    doc.add_paragraph("6- Anlaşmazlık halinde haksız olan taraf, bu sözleşmede yazılı tazminatlarla birlikte diğer tarafın maruz kalacağı her türlü zarar, ziyan, mahkeme ve icra masrafları ile vekalet ücretini de ödeyecektir.")
    doc.add_paragraph(f"7- Bu sözleşmenin uygulanmasından doğacak her türlü uyuşmazlığın giderilmesinde {mahkemeleri} Mahkemeleri ve icra daireleri yetkilidir  {date}")
    doc.add_paragraph('ALICI				EMLAK KOMİSYONCUSU				SATICI')

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

def rental_contract(tapu_dairesi, mahalle,sokak_cadde,bina_no,kiralananin_cinsi,
                    kiralayan_ad_soyad, kiralayan_tcno, kiralayan_adres,
                    kiraci_ad_soyad, kiraci_tcno, kiraci_adres, akdin_baslangic_tarihi,
                    akdin_suresi, yillik_kira_bedeli, aylik_kira_bedeli, kira_bedeli_odeme_sekli,
                    kiralanan_kullanim_sekli, kiralananin_durumu, demirbaslar, date):
    doc = docx.Document()
    heading = doc.add_heading('KİRA SÖZLEŞMESİ', 0)
    run = heading.runs[0]
    run.font.name = 'Time New Roman'
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.bold = True
    run.font.size = Pt(16)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'DAİRESİ    : {tapu_dairesi}')
    doc.add_paragraph(f'Mahallesi     : {mahalle}')
    doc.add_paragraph(f'Sokak/Cadde : {sokak_cadde}')
    doc.add_paragraph(f'Bina No - Daire No      : {bina_no}')
    doc.add_paragraph(f'Kiralananın Cinsi   : {kiralananin_cinsi}')
    doc.add_paragraph(f'Kiraya Veren   : {kiralayan_ad_soyad}')
    doc.add_paragraph(f'Kiraya Verenin T.C Kimlik No   : {kiralayan_tcno}')
    doc.add_paragraph(f'Kiraya Verenin Adresi  : {kiralayan_adres}')
    doc.add_paragraph(f'Kiracı  : {kiraci_ad_soyad}')
    doc.add_paragraph(f'Kiracının T.C Kimlik No : {kiraci_tcno}')
    doc.add_paragraph(f'Kiracının Adresi    : {kiraci_adres}')
    doc.add_paragraph(f'Akdinin Başlangıç Tarihi    : {akdin_baslangic_tarihi}')
    doc.add_paragraph(f'Akdin Süresi    : {akdin_suresi}')
    doc.add_paragraph(f'Yıllık Kira Bedeli  : {yillik_kira_bedeli}')
    doc.add_paragraph(f'Aylık Kira Bedeli   : {aylik_kira_bedeli}')
    doc.add_paragraph(f'Kira Bedelinin Ödeme Şekli  : {kira_bedeli_odeme_sekli}')
    doc.add_paragraph(f'Kiralananın Kullanım Şekli  : {kiralanan_kullanim_sekli}')
    doc.add_paragraph(f'Kiralananın Durumu  : {kiralananin_durumu}')
    doc.add_paragraph(f'Kiralanan ile Birlikte Teslim Edilen Demirbaşlar: {demirbaslar}')
    doc.add_paragraph('İMZA                 				İMZA                 				İMZA')

    doc.add_page_break()
    second_heading = doc.add_heading('KİRA SÖZLEŞMESİ ÖZEL HÜKÜMLERİ', 0)
    run = second_heading.runs[0]
    run.font.name = 'Time New Roman'
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.bold = True
    run.font.size = Pt(16)
    second_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('1- Kiracı kiraladığı şeyi kendi malı gibi kullanmaya ve bozulmasına ,evsaf meziyetlerini şöhret ve itibarını kaybetmesine meydan vermemeye mecburdur.')
    doc.add_paragraph('2- Kiralananın tahsis konut ise sadece aile bireyleri oturacaktır.Hiç bir şekilde üçüncü bir kişi ikamet eden sıfatıyla geçici veya daimi surette kiralanan da kalamaz ; Kiralanan işyeri ise burada sadece kiracı, kararlaştırılan konuda faaliyet yürütecektir.')
    doc.add_paragraph('3- Kiralanan yerin su,elektrik ,Doğalgaz ,yakıt masrafları, apartmanın aydınlatılması, temizlenmesi gibi nedenlerle doğacak apartman aidatları v.b giderler ile kapıcı parası kiracıya aittir.')
    doc.add_paragraph('4- Kiralanan malik tarafından satılığa çıkarılacak olursa ; kiracı müşteri adaylarının kiralananı gezip görmelerine müsaade edecektir')
    doc.add_paragraph('5- Kiracı kiraladığı şeyi ne halde buldu ise kiraya verene o halde teslim etmeye mecburdur. Ancak kiralananda, giderilmesi kiraya verenin sorumluluğu kapsamında bulunan herhangi bir arıza veya hasar meydana gelirse, kiracı durumu kiraya verene yazılı olarak ihbar edecektir.İhbar yapmadan ,kiracının kendiliğinden yapacağı harcamalardan kiraya veren sorumlu olmayacaktır. Kiracı kiraya verenin muvafakatı çerçevesinde anlaşarak, kiralananda bir takım faydalı tadilat,tamirat ve dekorasyon yapabilir.Tahliye sırasında kiraya verenin seçimlik hakkı yapılan anlaşma ile sınırlı olacaktır.')
    doc.add_paragraph('6- Kiracı, kiralanan gayrimenkulun kira bedeli üzerinden 193 Sayılı Gelir Vergisi Kanunun 94.maddesi gereğince % 20 oranında gelir vergisi tevkifatı yapacaklardır.')
    doc.add_paragraph('7- Kiralanan şeyin vergisi ve tamiri kiraya verene kullanılması için lazım gelen temizleme ıslah masrafları kiracıya aittir.Bu hususta âdete bakılır.')
    doc.add_paragraph('8- Kiracı bu kira kontratosunun ek 1ve 2 belirtilen hususi şartlar baki kalmak şartı ile,dönem sonunda kiralananı tahliye etmek isterlerse bu isteklerini dönem sonundan en az bir ay önce ,kiraya verene yazılı olarak bildirecektir.')
    doc.add_paragraph('9- Bu sözleşmede yazılı bulunmayan hükümlere ihtiyaç duyulduğunda 6570 sayılı Kira Kanunu,Medeni Kanun, Borçlar kanunu,634 Sayılı kat mülkiyeti kanunu ve diğer yürürlükteki alakalı kanun ve Yargıtay kararları uygulanır.')
    doc.add_paragraph('Tarafların özgür rızaları ile tanzim ve imza olunan ve 9 maddeden ibaret “ Özel hükümleri” içeren işbu sözleşme ; iki suret olarak düzenlenmiş ve taraflara birer sureti verilmiştir.')
    doc.add_paragraph(f'            			 Düzenleme Tarihi {date}')
    doc.add_paragraph('KİRACI									KİRAYA VEREN')

    stream=BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

tabs = st.tabs(["Kira Sözleşmesi", "Alım-Satım Sözleşmesi"])

with tabs[0]:
    st.header("Kira Sözleşmesi")
    with st.form("kira_form"):
        tapu_dairesi = st.text_input("Tapu Dairesi")
        mahalle = st.text_input("Mahalle")
        sokak_cadde = st.text_input('Sokak/Cadde')
        bina_no = st.text_input('Bina No / Daire No')
        kiralananin_cinsi = st.text_input("Kiralananın Cinsi")
        kiralayan_ad_soyad = st.text_input("Kiraya Veren Ad Soyad")
        kiralayan_tcno = st.text_input("Kiraya Veren T.C Kimlik No",max_chars=11)
        kiralayan_adres = st.text_area("Kiraya Veren Adres")
        kiraci_ad_soyad = st.text_input("Kiracı Ad Soyad")
        kiraci_tcno = st.text_input("Kiracının T.C Kimlik No",max_chars=11)
        kiraci_adres = st.text_area("Kiracı Adres")
        akdin_baslangic_tarihi = st.date_input("Başlangıç Tarihi",value=None)
        akdin_suresi = st.text_input("Akdin Süresi")
        yillik_kira_bedeli = st.text_input("Yıllık Kira Bedeli")
        aylik_kira_bedeli = st.text_input("Aylık Kira Bedeli")
        kira_bedeli_odeme_sekli = st.text_input("Kira Bedeli Ödeme Şekli")
        kiralanan_kullanim_sekli = st.text_input("Kiralanan Kullanım Şekli")
        kiralananin_durumu = st.text_input("Kiralananın Durumu")
        demirbaslar = st.text_area("Demirbaşlar")
        date = st.date_input("Sözleşme Tarihi",value=None)
        submitted = st.form_submit_button("Sözleşmeyi Oluştur")
    if submitted:
        stream = rental_contract(
            tapu_dairesi, mahalle,sokak_cadde,bina_no, kiralananin_cinsi,
            kiralayan_ad_soyad, kiralayan_tcno, kiralayan_adres,
            kiraci_ad_soyad, kiraci_tcno, kiraci_adres, akdin_baslangic_tarihi,
            akdin_suresi, yillik_kira_bedeli, aylik_kira_bedeli, kira_bedeli_odeme_sekli,
            kiralanan_kullanim_sekli, kiralananin_durumu, demirbaslar, date
        )
        st.download_button("Sözleşmeyi indir Word", stream.getvalue(), "kira_sozlesmesi.docx")

with tabs[1]:
    st.header("Alım-Satım Sözleşmesi")
    with st.form("alim_form"):
        cins = st.text_input("Cins")
        il = st.text_input("İl")
        ilce = st.text_input("İlçe")
        mahalle = st.text_input("Mahalle")
        sokak = st.text_input("Sokak")
        kapi_no = st.text_input("Kapı No")
        pafta_no = st.text_input("Pafta No")
        ada_no = st.text_input("Ada No")
        parsel_no = st.text_input("Parsel No")
        alici = st.text_input("Alıcı")
        satici = st.text_input("Satıcı")
        fiyat = st.text_input("Fiyat (TL)")
        kaparo = st.text_input("Kaparo (TL)")
        odeme_sekli = st.text_input("Ödeme Şekli")
        emlakci_komisyon_yuzdesi = st.text_input("Emlakçı Komisyon Yüzdesi")
        mahkemeleri = st.text_input("Mahkemeleri")
        date = st.text_input("Sözleşme Tarihi",placeholder="Sözleşmeyi yaptığınız tarihi giriniz.")
        submitted2 = st.form_submit_button("Sözleşmeyi Oluştur")
    if submitted2:
        stream = create_sales_contract(
            cins, il, ilce, mahalle, sokak, kapi_no,
            pafta_no, ada_no, parsel_no, alici, satici,
            fiyat, kaparo, odeme_sekli, emlakci_komisyon_yuzdesi,
            mahkemeleri,date
        )

        st.download_button("Sözleşmeyi indir Word", stream.getvalue(), "satis_sozlesmesi.docx")